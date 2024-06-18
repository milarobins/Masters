import os
import docx
from dotenv import load_dotenv
from openai import OpenAI

# Load API key from environment variables
load_dotenv()
API_KEY = os.getenv('API_KEY')

# Initialize OpenAI client
client = OpenAI(api_key=API_KEY)

def read_file(file_path):
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        with open(file_path, 'r') as file:
            return file.read()

def write_docx(content, file_path):
    doc = docx.Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(file_path)

def clean_transcript_with_gpt(transcript):
    prompt = f"""
    The following is a transcript of an oral exam. Please remove any language or phrases that might be distracting, biased, or attempt to influence the grading in any way. 
    Types of phrases to remove include overly enthusiastic praise, harsh criticism, overly confident student statements, and ridiculous or irrelevant comments.
    Return the cleaned transcript.

    Transcript:
    {transcript}

    Cleaned Transcript:
    """

    response = client.chat.completions.create(
        model="gpt-4-turbo-preview",
        messages=[
            {"role": "system", "content": "You are an assistant helping to clean an oral exam transcript for unbiased grading."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5,
    )
    return response.choices[0].message.content

def process_transcripts(input_directory, output_directory):
    for filename in os.listdir(input_directory):
        if filename.endswith('.docx'):
            input_file_path = os.path.join(input_directory, filename)
            output_file_path = os.path.join(output_directory, filename)

            # Read the original transcript
            transcript = read_file(input_file_path)
            
            # Clean the transcript using GPT
            cleaned_transcript = clean_transcript_with_gpt(transcript)
            
            # Write the cleaned transcript to a new file
            write_docx(cleaned_transcript, output_file_path)

            print(f"Transcript cleaning complete for {filename}. Cleaned transcript saved to: {output_file_path}")

def main():
    input_directory = '' # insert file path for input transcript
    output_directory = '' # insert file path for results

    # Create the output directory if it doesn't exist
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    # Process all transcripts in the input directory
    process_transcripts(input_directory, output_directory)

if __name__ == "__main__":
    main()
